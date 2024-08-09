#!/usr/bin/env python
# coding: utf-8
import pandas as pd
from docx import Document


class PlansSumWordFileGenerate:
    """
    用于生成段落格式为：
    “池州马衙站(舟山)  8月7日/ 7点-优卓
    车牌号：浙L29137/浙LK833挂
    驾驶员：李大成13946173662
    押运员：纪红霞13946050515
    承运商：舟山欣远运输有限公司（自有）” 的word文档；
    复思用。
    """

    def __init__(self, file):
        self.file = file

    def read_file(self):
        """read table"""
        try:
            df = pd.read_excel(self.file, header=1)
        except Exception as e:
            print(f"Error reading the Excel file: {e}")
            return None
        return df

    @staticmethod
    def reformat_date(date_str):
        """reformat python datetime to required string format"""
        month, day, time = date_str.split('-')
        return f"{int(month)}月{int(day)}日/ {int(time)}点"

    def date_formatting(self):
        """Execute the reformat_date function"""
        df = self.read_file()

        if df is not None:
            # df['计划日期'] = pd.to_datetime(df['计划日期'])
            df['计划日期'] = df['计划日期'].dt.strftime('%m-%d-%H')
            df['计划日期'] = df['计划日期'].apply(self.reformat_date)

        return df

    def generate_paragraphs(self):
        """generate paragraphs"""
        df = self.date_formatting()

        doc = Document()

        for index, row in df.iterrows():
            platform_name = row['到站名称']
            station_name = row['液源']
            planning_datetime = row['计划日期']
            customer_name = row['计划所属客户名称']
            car_number = row['车牌号/挂车号']
            driver_info = row['司机姓名/电话']
            supercargo_info = row['押运员姓名/电话']
            logistic_company = row['所属物流公司名称']

            formatted_paragraph = (f"{platform_name}（{station_name}） {planning_datetime}-{customer_name}\n"
                                   f"车牌号：{car_number}\n"
                                   f"驾驶员：{driver_info}\n"
                                   f"押运员：{supercargo_info}\n"
                                   f"承运商：{logistic_company}")

            doc.add_paragraph(formatted_paragraph)
        return doc

    def save_docx(self, save_path='test output.docx'):
        """download word file"""
        doc = self.generate_paragraphs()
        doc.save(save_path)


class TrucksWordFileGenerate:

    def __init__(self, filepath):
        self.filepath = filepath

    def load_dataframe(self):
        """read table"""
        try:
            df = pd.read_excel(self.filepath, dtype={'电话': str, '电话.1': str})
            df = df.dropna().reset_index(drop=True)
        except Exception as e:
            print(f"Error reading the Excel file: {e}")
            return None
        return df

    def generate_text(self):
        """generate paragraphs"""
        df = self.load_dataframe()

        doc = Document()

        for index, row in df.iterrows():
            car_number = row['车号']
            guache_number = row['挂车号']
            driver_name = row['驾驶员']
            driver_phone_number = row['电话']
            supercargo_name = row['押运员']
            supercargo_phone_number = row['电话.1']
            logistic_company = row['物流公司']

            formatted_paragraph = (f"车号：{car_number}\n"
                                   f"挂车号：{guache_number}\n"
                                   f"驾驶员：{driver_name} {driver_phone_number}\n"
                                   f"押运员：{supercargo_name} {supercargo_phone_number}\n"
                                   f"{logistic_company}")

            doc.add_paragraph(formatted_paragraph)
        return doc

    def save_docx(self, save_path='test output.docx'):
        """download file"""
        doc = self.generate_text()
        doc.save(save_path)


if __name__ == "__main__":
    trucks_info = TrucksWordFileGenerate('车辆信息.xlsx')
    trucks_info.save_docx('车辆信息.docx')

    logistic_info = PlansSumWordFileGenerate('计划汇集报送（暂定）.xlsx')
    logistic_info.save_docx('计划汇集报送（暂定）.docx')



# In[ ]:




